import streamlit as st
from docx import Document
import pandas as pd
import plotly.express as px
import google.generativeai as genai
import json
import os
import traceback
import re 

# --- 1. Configuração da Chave da API do Gemini ---
def get_gemini_api_key():
    try: return st.secrets["GOOGLE_API_KEY"]
    except (FileNotFoundError, KeyError): 
        api_key = os.environ.get("GOOGLE_API_KEY")
        return api_key if api_key else None

# --- 2. Funções de Processamento do Documento e Interação com Gemini ---
def parse_value_for_numeric(val_str_in):
    if pd.isna(val_str_in) or str(val_str_in).strip() == '': return None
    text = str(val_str_in).strip()
    is_negative_paren = text.startswith('(') and text.endswith(')')
    if is_negative_paren: text = text[1:-1]
    text_num_part = re.sub(r'[R$\s%]', '', text)
    if ',' in text_num_part and '.' in text_num_part:
        if text_num_part.rfind('.') < text_num_part.rfind(','): text_num_part = text_num_part.replace('.', '') 
        text_num_part = text_num_part.replace(',', '.') 
    elif ',' in text_num_part: text_num_part = text_num_part.replace(',', '.')
    match = re.search(r"([-+]?\d*\.?\d+|\d+)", text_num_part)
    if match:
        try: num = float(match.group(1)); return -num if is_negative_paren else num
        except ValueError: return None
    return None

def extrair_conteudo_docx(uploaded_file):
    try:
        document = Document(uploaded_file)
        textos = [p.text for p in document.paragraphs if p.text.strip()]
        tabelas_data = [] 
        for i, table_obj in enumerate(document.tables):
            data_rows, keys, nome_tabela = [], None, f"Tabela_DOCX_{i+1}"
            try: 
                prev_el = table_obj._element.getprevious()
                if prev_el is not None and prev_el.tag.endswith('p'):
                    p_text = "".join(node.text for node in prev_el.xpath('.//w:t')).strip()
                    if p_text and len(p_text) < 80: nome_tabela = p_text.replace(":", "").strip()
            except: pass
            for r_idx, row in enumerate(table_obj.rows):
                cells = [c.text.strip() for c in row.cells]
                if r_idx == 0: keys = [k.replace("\n"," ").strip() if k else f"Col{c_idx+1}" for c_idx, k in enumerate(cells)]; continue
                if keys: 
                    row_dict = {}
                    for k_idx, key_name in enumerate(keys):
                        row_dict[key_name] = cells[k_idx] if k_idx < len(cells) else None
                    data_rows.append(row_dict)
            if data_rows:
                try:
                    df = pd.DataFrame(data_rows)
                    for col in df.columns:
                        original_series = df[col].copy()
                        num_series = original_series.astype(str).apply(parse_value_for_numeric)
                        if num_series.notna().sum() / max(1, len(num_series)) > 0.3:
                            df[col] = pd.to_numeric(num_series, errors='coerce')
                            continue 
                        else: df[col] = original_series 
                        try:
                            temp_str_col = df[col].astype(str)
                            dt_series = pd.to_datetime(temp_str_col, errors='coerce', dayfirst=True) 
                            if dt_series.notna().sum() / max(1, len(dt_series)) > 0.5:
                                df[col] = dt_series
                            else: 
                                df[col] = original_series.astype(str).fillna('')
                        except: df[col] = original_series.astype(str).fillna('')
                    for col in df.columns: 
                        if df[col].dtype == 'object': df[col] = df[col].astype(str).fillna('')
                    tabelas_data.append({"id": f"doc_tabela_{i+1}", "nome": nome_tabela, "dataframe": df})
                except Exception as e_df_proc:
                    st.warning(f"Não foi possível processar DataFrame para tabela '{nome_tabela}': {e_df_proc}")
        return "\n\n".join(textos), tabelas_data
    except Exception as e_doc_read: 
        st.error(f"Erro crítico ao ler DOCX: {e_doc_read}")
        return "", []

# --- Sistema Multi-Agente para Análise Profunda ---
def agente_analise_dados(texto_doc, tabelas_info_list, model):
    """Agente especializado em análise de dados e métricas"""
    tabelas_prompt_str = ""
    for t_info in tabelas_info_list:
        df, nome_t, id_t = t_info["dataframe"], t_info["nome"], t_info["id"]
        sample_df = df.head(5).iloc[:, :min(8, len(df.columns))]
        try: md_table = sample_df.to_markdown(index=False)
        except: md_table = sample_df.to_string(index=False) 
        
        # Análise estatística básica
        stats_info = ""
        for col in df.columns[:5]:  # Primeiras 5 colunas
            if pd.api.types.is_numeric_dtype(df[col]):
                stats = df[col].describe()
                stats_info += f"\n  {col}: Média={stats['mean']:.2f}, Mediana={stats['50%']:.2f}, Min={stats['min']:.2f}, Max={stats['max']:.2f}"
        
        tabelas_prompt_str += f"\n--- Tabela '{nome_t}' (ID: {id_t}) ---\nEstatísticas:{stats_info}\nDados:\n{md_table}\n"
    
    text_limit = 20000
    prompt_text = texto_doc[:text_limit]
    
    prompt = f"""
    Você é um AGENTE ESPECIALISTA EM ANÁLISE DE DADOS. Analise profundamente os dados fornecidos.
    
    [TEXTO]{prompt_text}[FIM TEXTO]
    [TABELAS]{tabelas_prompt_str}[FIM TABELAS]
    
    Como especialista em dados, identifique:
    1. KPIs críticos e métricas importantes
    2. Tendências e padrões nos dados
    3. Correlações entre variáveis
    4. Outliers ou anomalias
    5. Insights quantitativos profundos
    
    Retorne JSON com:
    {{
        "kpis_criticos": [
            {{"nome": "Nome KPI", "valor": "Valor", "insight": "Insight sobre o KPI", "criticidade": "alta|media|baixa"}}
        ],
        "tendencias": ["Tendência 1", "Tendência 2"],
        "correlacoes": ["Correlação identificada 1", "Correlação 2"],
        "anomalias": ["Anomalia 1", "Anomalia 2"],
        "insights_quantitativos": ["Insight 1", "Insight 2"]
    }}
    """
    
    try:
        response = model.generate_content(prompt)
        return json.loads(response.text.strip().lstrip("```json").rstrip("```").strip())
    except: return {}

def agente_analise_estrategica(texto_doc, tabelas_info_list, model):
    """Agente especializado em análise estratégica e de negócios"""
    text_limit = 25000
    prompt_text = texto_doc[:text_limit]
    
    prompt = f"""
    Você é um AGENTE ESPECIALISTA EM ANÁLISE ESTRATÉGICA E DE NEGÓCIOS. Faça uma análise estratégica profunda.
    
    [TEXTO]{prompt_text}[FIM TEXTO]
    
    Como especialista estratégico, analise:
    1. Pontos fortes e fracos estratégicos
    2. Oportunidades de mercado e crescimento
    3. Ameaças e riscos
    4. Recomendações acionáveis
    5. Cenários futuros possíveis
    
    Retorne JSON com:
    {{
        "analise_swot_detalhada": {{
            "forcas": ["Força detalhada 1", "Força 2"],
            "fraquezas": ["Fraqueza detalhada 1", "Fraqueza 2"],
            "oportunidades": ["Oportunidade detalhada 1", "Oportunidade 2"],
            "ameacas": ["Ameaça detalhada 1", "Ameaça 2"]
        }},
        "recomendacoes_acionaveis": [
            {{"acao": "Ação recomendada", "prioridade": "alta|media|baixa", "prazo": "curto|medio|longo", "impacto": "Descrição do impacto"}}
        ],
        "cenarios_futuros": ["Cenário 1", "Cenário 2"],
        "fatores_criticos_sucesso": ["Fator 1", "Fator 2"]
    }}
    """
    
    try:
        response = model.generate_content(prompt)
        return json.loads(response.text.strip().lstrip("```json").rstrip("```").strip())
    except: return {}

def agente_sintese(analise_dados, analise_estrategica, texto_doc, model):
    """Agente que sintetiza insights de outros agentes"""
    
    prompt = f"""
    Você é um AGENTE SINTETIZADOR SÊNIOR. Combine e sintetize os insights dos agentes especialistas.
    
    ANÁLISE DE DADOS: {json.dumps(analise_dados, ensure_ascii=False)}
    ANÁLISE ESTRATÉGICA: {json.dumps(analise_estrategica, ensure_ascii=False)}
    
    Como sintetizador sênior, crie:
    1. Síntese executiva dos principais insights
    2. Conexões entre insights quantitativos e estratégicos
    3. Priorização de ações baseada em dados
    4. Roadmap de implementação
    
    Retorne JSON com:
    {{
        "sintese_executiva": "Resumo executivo dos principais insights...",
        "conexoes_insights": ["Conexão 1 entre dados e estratégia", "Conexão 2"],
        "priorizacao_acoes": [
            {{"acao": "Ação prioritária", "score": 1-10, "justificativa": "Por que é prioritária"}}
        ],
        "roadmap": {{
            "imediato": ["Ação imediata 1", "Ação 2"],
            "30_dias": ["Ação 30 dias 1", "Ação 2"],
            "90_dias": ["Ação 90 dias 1", "Ação 2"]
        }}
    }}
    """
    
    try:
        response = model.generate_content(prompt)
        return json.loads(response.text.strip().lstrip("```json").rstrip("```").strip())
    except: return {}

def analisar_documento_com_gemini(texto_doc, tabelas_info_list):
    """Análise original com visualizações"""
    api_key = get_gemini_api_key()
    if not api_key: st.warning("Chave API Gemini não configurada."); return []
    try:
        genai.configure(api_key=api_key)
        safety_settings = [{"category": c,"threshold": "BLOCK_NONE"} for c in ["HARM_CATEGORY_HARASSMENT","HARM_CATEGORY_HATE_SPEECH","HARM_CATEGORY_SEXUALLY_EXPLICIT","HARM_CATEGORY_DANGEROUS_CONTENT"]]
        model = genai.GenerativeModel(model_name="gemini-1.5-flash-latest", safety_settings=safety_settings)
        
        tabelas_prompt_str = ""
        for t_info in tabelas_info_list:
            df, nome_t, id_t = t_info["dataframe"], t_info["nome"], t_info["id"]
            sample_df = df.head(3).iloc[:, :min(5, len(df.columns))]
            md_table = ""
            try: md_table = sample_df.to_markdown(index=False)
            except: md_table = sample_df.to_string(index=False) 
            
            colunas_para_mostrar_tipos = df.columns.tolist()[:min(8, len(df.columns))]
            col_types_list = [f"'{col_name_prompt}' (tipo: {str(df[col_name_prompt].dtype)})" for col_name_prompt in colunas_para_mostrar_tipos]
            col_types_str = ", ".join(col_types_list)
            
            tabelas_prompt_str += f"\n--- Tabela '{nome_t}' (ID: {id_t}) ---\nColunas e tipos (primeiras {len(colunas_para_mostrar_tipos)}): {col_types_str}\nAmostra de dados:\n{md_table}\n"
        
        text_limit = 45000
        prompt_text = texto_doc[:text_limit] + ("\n[TEXTO TRUNCADO...]" if len(texto_doc) > text_limit else "")
        
        prompt = f"""
        Você é um assistente de análise de dados. Analise o texto e as tabelas.
        [TEXTO]{prompt_text}[FIM TEXTO]
        [TABELAS]{tabelas_prompt_str}[FIM TABELAS]

        Gere lista JSON de sugestões de visualizações. Objeto DEVE ter: "id", "titulo", "tipo_sugerido" ("kpi", "tabela_dados", "lista_swot", "grafico_barras", "grafico_pizza", "grafico_linha", "grafico_dispersao", "grafico_barras_agrupadas"), "fonte_id" (ID tabela ou "texto_descricao_fonte"), "parametros" (objeto JSON), "justificativa".
        Para "parametros":
        - "kpi": {{"valor": "ValorKPI", "delta": "Mudança", "descricao": "Contexto"}}
        - "tabela_dados": Para TABELA EXISTENTE: {{"id_tabela_original": "ID_Tabela"}}. Para DADOS DO TEXTO: {{"dados": [{{"Coluna1": "ValorA1"}}, ...], "colunas_titulo": ["Título Col1"]}}
        - "lista_swot": {{"forcas": ["F1"], "fraquezas": ["Fr1"], "oportunidades": ["Op1"], "ameacas": ["Am1"]}} (Listas de strings).
        - Gráficos de TABELA ("barras", "linha", "dispersao"): {{"eixo_x": "NOME_COL_X", "eixo_y": "NOME_COL_Y"}} (Y numérico).
        - Gráficos de PIZZA de TABELA: {{"categorias": "NOME_COL_CAT", "valores": "NOME_COL_VAL_NUM"}} (Valores numéricos).
        - Gráficos com DADOS EXTRAÍDOS DO TEXTO ("barras", "pizza", etc.): {{"dados": [{{"NomeEixoX": "CatA", "NomeEixoY": ValNumA}}, ...], "eixo_x": "NomeEixoX", "eixo_y": "NomeEixoY"}} (Valores DEVEM ser numéricos).
        - "grafico_barras_agrupadas": Se de TABELA: {{"eixo_x": "COL_PRINCIPAL", "eixo_y": "COL_VALOR_NUM", "cor_agrupamento": "COL_SUB_CAT"}}. Se DADOS EXTRAÍDOS: {{"dados": [{{"CatPrincipal": "A", "SubCat": "X", "Valor": 10}}, ...], "eixo_x": "CatPrincipal", "eixo_y": "Valor", "cor_agrupamento": "SubCat"}}.
        
        INSTRUÇÕES CRÍTICAS:
        1.  NOMES DE COLUNAS: Para gráficos de TABELA, use os NOMES EXATOS das colunas.
        2.  DADOS NUMÉRICOS: Se coluna de valor de TABELA não for numérica, NÃO sugira gráfico que precise de número para ela, A MENOS que extraia valor numérico dela (ex: '70%' -> 70.0; '70% - 86%' -> 70.0). Se extrair do texto, coloque em "dados", garanta valores numéricos.
        3.  COBERTURA GEOGRÁFICA (Player, Cidades): Se lista, sugira "tabela_dados" com "dados" nos "parametros" e "colunas_titulo". Não "mapa".
        4.  SWOT: Se tabela compara SWOTs, gere "lista_swot" INDIVIDUAL por player.
        Retorne APENAS a lista JSON válida.
        """
        with st.spinner("🤖 Gemini analisando visualizações..."):
            response = model.generate_content(prompt)
        cleaned_text = response.text.strip().lstrip("```json").rstrip("```").strip()
        sugestoes = json.loads(cleaned_text)
        if isinstance(sugestoes, list) and all(isinstance(item, dict) for item in sugestoes): st.success(f"{len(sugestoes)} sugestões!"); return sugestoes
        st.error("Resposta Gemini não é lista JSON."); return []
    except json.JSONDecodeError as e: st.error(f"Erro JSON Gemini: {e}"); st.code(response.text if 'response' in locals() else "N/A", language="text"); return []
    except Exception as e: st.error(f"Erro API Gemini: {e}"); st.text(traceback.format_exc()); return []

def executar_analise_profunda_multiagente(texto_doc, tabelas_info_list):
    """Executa análise profunda usando sistema multi-agente"""
    api_key = get_gemini_api_key()
    if not api_key: 
        st.warning("Chave API Gemini não configurada para análise profunda.")
        return {}
    
    try:
        genai.configure(api_key=api_key)
        safety_settings = [{"category": c,"threshold": "BLOCK_NONE"} for c in ["HARM_CATEGORY_HARASSMENT","HARM_CATEGORY_HATE_SPEECH","HARM_CATEGORY_SEXUALLY_EXPLICIT","HARM_CATEGORY_DANGEROUS_CONTENT"]]
        model = genai.GenerativeModel(model_name="gemini-1.5-flash-latest", safety_settings=safety_settings)
        
        with st.spinner("🔍 Executando análise profunda multi-agente..."):
            # Executar agentes em paralelo conceitual
            col1, col2 = st.columns(2)
            
            with col1:
                st.info("🔢 Agente de Análise de Dados trabalhando...")
                try:
                    analise_dados = agente_analise_dados(texto_doc, tabelas_info_list, model)
                    if analise_dados:
                        st.success("✅ Análise de dados concluída")
                    else:
                        st.warning("⚠️ Análise de dados parcial")
                        analise_dados = {}
                except Exception as e:
                    st.error(f"❌ Erro na análise de dados: {str(e)[:100]}...")
                    analise_dados = {}
            
            with col2:
                st.info("📊 Agente de Análise Estratégica trabalhando...")
                try:
                    analise_estrategica = agente_analise_estrategica(texto_doc, tabelas_info_list, model)
                    if analise_estrategica:
                        st.success("✅ Análise estratégica concluída")
                    else:
                        st.warning("⚠️ Análise estratégica parcial")
                        analise_estrategica = {}
                except Exception as e:
                    st.error(f"❌ Erro na análise estratégica: {str(e)[:100]}...")
                    analise_estrategica = {}
            
            st.info("🧠 Agente Sintetizador integrando insights...")
            try:
                sintese = agente_sintese(analise_dados, analise_estrategica, texto_doc, model)
                if sintese:
                    st.success("✅ Síntese concluída - Análise profunda finalizada!")
                else:
                    st.warning("⚠️ Síntese parcial")
                    sintese = {}
            except Exception as e:
                st.error(f"❌ Erro na síntese: {str(e)[:100]}...")
                sintese = {}
        
        # Verificar se pelo menos um agente funcionou
        if not analise_dados and not analise_estrategica and not sintese:
            st.error("Não foi possível executar a análise profunda. Tente novamente ou verifique a configuração da API.")
            return {}
        
        return {
            "analise_dados": analise_dados,
            "analise_estrategica": analise_estrategica,
            "sintese": sintese,
            "status": "sucesso_parcial" if not all([analise_dados, analise_estrategica, sintese]) else "sucesso_completo"
        }
    
    except Exception as e:
        st.error(f"Erro na análise profunda multi-agente: {e}")
        return {}

def render_kpis(kpi_sugestoes):
    if kpi_sugestoes:
        num_kpis = len(kpi_sugestoes); kpi_cols = st.columns(min(num_kpis, 4)) 
        for i, kpi_sug in enumerate(kpi_sugestoes):
            with kpi_cols[i % min(num_kpis, 4)]:
                params=kpi_sug.get("parametros",{}); delta_val=str(params.get("delta",""))
                st.metric(label=kpi_sug.get("titulo","KPI"),value=str(params.get("valor","N/A")),delta=delta_val if delta_val else None,help=params.get("descricao"))
        st.divider()

def render_swot_card(titulo_completo_swot, swot_data, card_key_prefix=""):
    st.subheader(f"{titulo_completo_swot}") 
    col1, col2 = st.columns(2)
    swot_map = {"forcas": ("Forças 💪", col1), "fraquezas": ("Fraquezas 📉", col1), 
                "oportunidades": ("Oportunidades 🚀", col2), "ameacas": ("Ameaças ⚠️", col2)}
    for key_swot_category, (header_swot_render, col_target_swot_render) in swot_map.items():
        with col_target_swot_render:
            st.markdown(f"##### {header_swot_render}")
            points_swot_render = swot_data.get(key_swot_category, ["N/A (info. não fornecida)"])
            if not points_swot_render or not isinstance(points_swot_render, list) or not all(isinstance(p_swot, str) for p_swot in points_swot_render): 
                points_swot_render = ["N/A (formato de dados incorreto)"]
            if not points_swot_render: points_swot_render = ["N/A"] 
            for item_swot_render in points_swot_render: 
                st.markdown(f"<div style='margin-bottom: 5px;'>- {item_swot_render}</div>", unsafe_allow_html=True)
    st.markdown("<hr style='margin-top: 10px; margin-bottom: 20px;'>", unsafe_allow_html=True)

def render_plotly_chart(item_config, df_plot_input):
    if df_plot_input is None:
        st.warning(f"Dados não disponíveis para o gráfico '{item_config.get('titulo', 'Sem Título')}'.")
        return False
    df_plot = df_plot_input.copy()
    tipo_grafico, titulo, params = item_config.get("tipo_sugerido"), item_config.get("titulo"), item_config.get("parametros", {})
    x_col, y_col, cat_col, val_col = params.get("eixo_x"), params.get("eixo_y"), params.get("categorias"), params.get("valores")
    cor_agrupamento_col = params.get("cor_agrupamento")
    fig, plot_func, plot_args = None, None, {}
    if tipo_grafico in ["grafico_barras", "grafico_barras_agrupadas"] and x_col and y_col: 
        plot_func,plot_args=px.bar,{"x":x_col,"y":y_col}
        if tipo_grafico == "grafico_barras_agrupadas" and cor_agrupamento_col:
            plot_args["color"], plot_args["barmode"] = cor_agrupamento_col, "group"
    elif tipo_grafico=="grafico_linha" and x_col and y_col: plot_func,plot_args=px.line,{"x":x_col,"y":y_col,"markers":True}
    elif tipo_grafico=="grafico_dispersao" and x_col and y_col: plot_func,plot_args=px.scatter,{"x":x_col,"y":y_col}
    elif tipo_grafico=="grafico_pizza" and cat_col and val_col: plot_func,plot_args=px.pie,{"names":cat_col,"values":val_col}
    if plot_func:
        required_cols=[col for col in plot_args.values() if isinstance(col,str)]
        if not all(col in df_plot.columns for col in required_cols):
            st.warning(f"Colunas {required_cols} não encontradas para '{titulo}'. Disponíveis: {df_plot.columns.tolist()}")
            return False
        try:
            df_plot_cleaned = df_plot.copy() 
            y_ax,val_ax=plot_args.get("y"),plot_args.get("values")
            if y_ax and y_ax in df_plot_cleaned.columns: df_plot_cleaned[y_ax]=pd.to_numeric(df_plot_cleaned[y_ax],errors='coerce')
            if val_ax and val_ax in df_plot_cleaned.columns: df_plot_cleaned[val_ax]=pd.to_numeric(df_plot_cleaned[val_ax],errors='coerce')
            cols_to_check_na_final = [val_check for val_check in plot_args.values() if isinstance(val_check, str) and val_check in df_plot_cleaned.columns]
            df_plot_cleaned.dropna(subset=cols_to_check_na_final, inplace=True)
            if not df_plot_cleaned.empty:
                fig=plot_func(df_plot_cleaned,title=titulo,**plot_args); st.plotly_chart(fig,use_container_width=True); return True
            else: st.warning(f"Dados insuficientes para '{titulo}' após limpar NaNs de {cols_to_check_na_final}.")
        except Exception as e_plotly_render: 
            st.warning(f"Erro ao gerar gráfico Plotly '{titulo}': {e_plotly_render}. Dtypes: {df_plot.dtypes.to_dict() if df_plot is not None else 'DF é None'}")
    elif tipo_grafico in ["grafico_barras","grafico_barras_agrupadas","grafico_linha","grafico_dispersao","grafico_pizza","grafico_radar"]:
        st.warning(f"Configuração de parâmetros incompleta para '{titulo}' (tipo: {tipo_grafico}).")
    return False

# --- 3. Interface Streamlit Principal ---
st.set_page_config(layout="wide", page_title="Gemini DOCX Insights")
for k, dv in [("s_gemini",[]),("cfg_sugs",{}),("doc_ctx",{"texto":"","tabelas":[]}),
              ("f_name",None),("dbg_cb_key",False),("pg_sel","Dashboard Principal"),("analise_profunda",{})]:
    st.session_state.setdefault(k, dv)

st.sidebar.title("✨ Navegação"); pg_opts_sb = ["Dashboard Principal","Análise SWOT Detalhada","Análise Profunda Multi-Agente"]
st.session_state.pg_sel=st.sidebar.radio("Selecione:",pg_opts_sb,index=pg_opts_sb.index(st.session_state.pg_sel),key="nav_radio_final_v7")
st.sidebar.divider(); uploaded_file_sb = st.sidebar.file_uploader("Selecione DOCX",type="docx",key="uploader_sidebar_final_v7")
# A chave do widget 'debug_cb_sidebar_key_final_v7' atualiza st.session_state.dbg_cb_key
st.session_state.dbg_cb_key=st.sidebar.checkbox("Mostrar Debug Info",value=st.session_state.dbg_cb_key,key="debug_cb_sidebar_final_v7")

if uploaded_file_sb:
    if st.session_state.f_name!=uploaded_file_sb.name: 
        with st.spinner("Processando novo documento..."):
            st.session_state.s_gemini,st.session_state.cfg_sugs,st.session_state.analise_profunda=[],{},{}
            st.session_state.f_name=uploaded_file_sb.name
            txt_main,tbls_main=extrair_conteudo_docx(uploaded_file_sb);st.session_state.doc_ctx={"texto":txt_main,"tabelas":tbls_main}
            if txt_main or tbls_main:
                sugs_main=analisar_documento_com_gemini(txt_main,tbls_main);st.session_state.s_gemini=sugs_main
                st.session_state.cfg_sugs={s.get("id",f"s_main_{i}_{hash(s.get('titulo'))}"):{"aceito":True,"titulo_editado":s.get("titulo","S/T"),"dados_originais":s} for i,s in enumerate(sugs_main)}
                
                # Executar análise profunda multi-agente
                st.session_state.analise_profunda = executar_analise_profunda_multiagente(txt_main, tbls_main)
            else: st.sidebar.warning("Nenhum conteúdo extraído.")
    
    if st.session_state.dbg_cb_key and (st.session_state.doc_ctx["texto"] or st.session_state.doc_ctx["tabelas"]): # Usa o estado correto
        with st.expander("Debug: Conteúdo DOCX (após extração e tipos)",expanded=False):
            st.text_area("Texto (amostra)",st.session_state.doc_ctx["texto"][:1000],height=80)
            for t_dbg_main in st.session_state.doc_ctx["tabelas"]:
                st.write(f"ID: {t_dbg_main['id']}, Nome: {t_dbg_main['nome']}")
                try: st.dataframe(t_dbg_main['dataframe'].head().astype(str).fillna("-"))
                except: st.text(f"Head:\n{t_dbg_main['dataframe'].head().to_string(na_rep='-')}")
                st.write("Tipos:",t_dbg_main['dataframe'].dtypes.to_dict())

    if st.session_state.s_gemini:
        st.sidebar.divider();st.sidebar.header("⚙️ Configurar Sugestões")
        for sug_cfg_loop in st.session_state.s_gemini:
            s_id_loop = sug_cfg_loop.get('id') 
            if not s_id_loop : continue 
            if s_id_loop not in st.session_state.cfg_sugs:
                 st.session_state.cfg_sugs[s_id_loop]={"aceito":True,"titulo_editado":sug_cfg_loop.get("titulo","S/T"),"dados_originais":sug_cfg_loop}
            cfg_loop = st.session_state.cfg_sugs[s_id_loop]
            
            with st.sidebar.expander(f"{cfg_loop['titulo_editado']}",expanded=False):
                st.caption(f"Tipo: {sug_cfg_loop.get('tipo_sugerido')} | Fonte: {sug_cfg_loop.get('fonte_id')}")
                cfg_loop["aceito"]=st.checkbox("Incluir?",value=cfg_loop["aceito"],key=f"acc_loop_{s_id_loop}")
                cfg_loop["titulo_editado"]=st.text_input("Título",value=cfg_loop["titulo_editado"],key=f"tit_loop_{s_id_loop}")
else: 
    if st.session_state.pg_sel=="Dashboard Principal": st.info("Upload DOCX na barra lateral.")

if st.session_state.pg_sel=="Dashboard Principal":
    st.title("📊 Dashboard de Insights")
    
    # Mostrar indicador de análise profunda disponível
    if uploaded_file_sb and st.session_state.analise_profunda:
        st.success("🧠 **Análise Profunda Multi-Agente** disponível! Acesse na barra lateral para insights mais detalhados.")
        st.divider()
    
    if uploaded_file_sb and st.session_state.s_gemini:
        kpis_r, outros_r = [], []
        for s_id_main_dash, s_cfg_main_dash in st.session_state.cfg_sugs.items():
            if s_cfg_main_dash["aceito"]: 
                item_main_dash = {"titulo":s_cfg_main_dash["titulo_editado"], **s_cfg_main_dash["dados_originais"]}
                (kpis_r if item_main_dash.get("tipo_sugerido")=="kpi" else outros_r).append(item_main_dash)
        
        render_kpis(kpis_r)
        
        if st.session_state.dbg_cb_key: # Usa o estado correto
             with st.expander("Debug: Elementos para Dashboard Principal (Não-KPI)",expanded=True): 
                st.json({"Outros Elementos (Configurados e Aceitos)": outros_r}, expanded=False)
        
        elementos_renderizados_dash = 0 
        col_idx_dash = 0 
        if outros_r:
            item_cols_main_dash = st.columns(2)
            for item_render_loop in outros_r:
                if item_render_loop.get("tipo_sugerido") == "lista_swot": continue 
                
                el_rend_d = False 
                with item_cols_main_dash[col_idx_dash % 2]: 
                    df_plot_loop = None 
                    params_loop = item_render_loop.get("parametros",{})
                    tipo_loop = item_render_loop.get("tipo_sugerido")
                    fonte_loop = item_render_loop.get("fonte_id")
                    titulo_loop = item_render_loop.get("titulo", "Visualização")
                    
                    st.subheader(titulo_loop) 
                    try:
                        if params_loop.get("dados"): 
                            try: df_plot_loop=pd.DataFrame(params_loop["dados"])
                            except Exception as e_dfd_loop: st.warning(f"'{titulo_loop}': Erro DF 'dados': {e_dfd_loop}"); continue
                        elif str(fonte_loop).startswith("doc_tabela_"): 
                            df_plot_loop=next((t["dataframe"] for t in st.session_state.doc_ctx["tabelas"] if t["id"]==fonte_loop),None)
                        
                        if tipo_loop=="tabela_dados":
                            df_tbl_loop=None
                            if str(fonte_loop).startswith("texto_") and params_loop.get("dados"):
                                try: 
                                    df_tbl_loop=pd.DataFrame(params_loop.get("dados")); 
                                    if params_loop.get("colunas_titulo"):df_tbl_loop.columns=params_loop.get("colunas_titulo")
                                except Exception as e_dftxt_loop: st.warning(f"Erro tabela texto '{titulo_loop}': {e_dftxt_loop}")
                            else: 
                                id_tbl_loop=params_loop.get("id_tabela_original",fonte_loop)
                                df_tbl_loop=next((t["dataframe"] for t in st.session_state.doc_ctx["tabelas"] if t["id"]==id_tbl_loop),None)
                            
                            if df_tbl_loop is not None: 
                                try: st.dataframe(df_tbl_loop.astype(str).fillna("-"))
                                except: st.text(df_tbl_loop.to_string(na_rep='-')); 
                                el_rend_d=True
                            else: st.warning(f"Tabela '{titulo_loop}' (Fonte: {fonte_loop}) não encontrada.")
                        
                        elif tipo_loop in ["grafico_barras","grafico_linha","grafico_dispersao","grafico_pizza", "grafico_barras_agrupadas"]:
                            if render_plotly_chart(item_render_loop, df_plot_loop): el_rend_d = True
                        
                        elif tipo_loop == 'mapa': 
                            st.info(f"Mapa para '{titulo_loop}' não implementado.")
                            el_rend_d=True
                        
                        if not el_rend_d and tipo_loop not in ["kpi","lista_swot","mapa"]: 
                            st.info(f"'{titulo_loop}' ({tipo_loop}) não gerado. Dados/Tipo não suportado.")
                    except Exception as e_render_loop: 
                        st.error(f"Erro renderizando '{titulo_loop}': {e_render_loop}")
                
                if el_rend_d: 
                    col_idx_dash += 1 
                    elementos_renderizados_dash += 1 
            
            if elementos_renderizados_dash == 0 and any(c['aceito'] and c['dados_originais'].get('tipo_sugerido') not in ['kpi','lista_swot'] for c in st.session_state.cfg_sugs.values()):
                st.info("Nenhum gráfico/tabela (além de KPIs/SWOTs) pôde ser gerado para o Dashboard Principal.")
        
        elif not kpis_r and not uploaded_file_sb: pass 
        elif not kpis_r and not outros_r and uploaded_file_sb and st.session_state.s_gemini: 
            st.info("Nenhum elemento selecionado ou gerado para o dashboard principal.")

elif st.session_state.pg_sel=="Análise SWOT Detalhada":
    st.title("🔬 Análise SWOT Detalhada")
    if not uploaded_file_sb: st.warning("Upload DOCX na barra lateral.")
    elif not st.session_state.s_gemini: st.info("Aguardando processamento/sugestões.")
    else:
        swot_sugs_page_render=[s_cfg_swot["dados_originais"] for s_id_swot,s_cfg_swot in st.session_state.cfg_sugs.items() if s_cfg_swot["aceito"] and s_cfg_swot["dados_originais"].get("tipo_sugerido")=="lista_swot"]
        if not swot_sugs_page_render: st.info("Nenhuma análise SWOT sugerida/selecionada.")
        else:
            if st.session_state.dbg_cb_key: # Usa o estado correto
                with st.expander("Debug: Dados para Análise SWOT (Página Dedicada)",expanded=False):st.json({"SWOTs Selecionados":swot_sugs_page_render})
            for swot_item_render_page in swot_sugs_page_render:
                render_swot_card(swot_item_render_page.get("titulo","SWOT"),swot_item_render_page.get("parametros",{}), card_key_prefix=swot_item_render_page.get("id","swot_pg_def"))

elif st.session_state.pg_sel=="Análise Profunda Multi-Agente":
    st.title("🧠 Análise Profunda Multi-Agente")
    if not uploaded_file_sb: 
        st.warning("Upload DOCX na barra lateral.")
    elif not st.session_state.analise_profunda:
        st.info("Aguardando análise profunda...")
    else:
        analise = st.session_state.analise_profunda
        
        # Mostrar status da análise
        status = analise.get("status", "desconhecido")
        if status == "sucesso_completo":
            st.success("✅ **Análise Profunda Completa** - Todos os agentes executaram com sucesso!")
        elif status == "sucesso_parcial":
            st.warning("⚠️ **Análise Profunda Parcial** - Alguns agentes apresentaram limitações, mas há insights disponíveis.")
        else:
            st.info("ℹ️ **Análise Profunda Disponível** - Visualize os insights gerados abaixo.")
        
        st.divider()
        
        # Síntese Executiva
        if analise.get("sintese", {}).get("sintese_executiva"):
            st.header("📋 Síntese Executiva")
            st.info(analise["sintese"]["sintese_executiva"])
            st.divider()
        
        # Análise de Dados
        if analise.get("analise_dados"):
            st.header("🔢 Insights de Dados")
            dados = analise["analise_dados"]
            
            col1, col2 = st.columns(2)
            
            with col1:
                if dados.get("kpis_criticos"):
                    st.subheader("📊 KPIs Críticos")
                    for kpi in dados["kpis_criticos"]:
                        criticidade_color = {"alta": "🔴", "media": "🟡", "baixa": "🟢"}.get(kpi.get("criticidade", "media"), "🟡")
                        st.metric(
                            label=f"{criticidade_color} {kpi.get('nome', 'KPI')}",
                            value=kpi.get('valor', 'N/A'),
                            help=kpi.get('insight', '')
                        )
                
                if dados.get("tendencias"):
                    st.subheader("📈 Tendências Identificadas")
                    for trend in dados["tendencias"]:
                        st.markdown(f"• {trend}")
            
            with col2:
                if dados.get("correlacoes"):
                    st.subheader("🔗 Correlações")
                    for corr in dados["correlacoes"]:
                        st.markdown(f"• {corr}")
                
                if dados.get("anomalias"):
                    st.subheader("⚠️ Anomalias Detectadas")
                    for anom in dados["anomalias"]:
                        st.warning(f"• {anom}")
            
            if dados.get("insights_quantitativos"):
                st.subheader("🧮 Insights Quantitativos")
                for insight in dados["insights_quantitativos"]:
                    st.markdown(f"• {insight}")
            
            st.divider()
        
        # Análise Estratégica
        if analise.get("analise_estrategica"):
            st.header("🎯 Análise Estratégica")
            estrategica = analise["analise_estrategica"]
            
            # SWOT Detalhada
            if estrategica.get("analise_swot_detalhada"):
                swot = estrategica["analise_swot_detalhada"]
                render_swot_card("SWOT Estratégico Detalhado", swot, "analise_profunda_swot")
            
            # Recomendações Acionáveis
            if estrategica.get("recomendacoes_acionaveis"):
                st.subheader("🎯 Recomendações Acionáveis")
                for rec in estrategica["recomendacoes_acionaveis"]:
                    prioridade_color = {"alta": "🔴", "media": "🟡", "baixa": "🟢"}.get(rec.get("prioridade", "media"), "🟡")
                    prazo_icon = {"curto": "⚡", "medio": "⏳", "longo": "🕐"}.get(rec.get("prazo", "medio"), "⏳")
                    
                    with st.expander(f"{prioridade_color} {prazo_icon} {rec.get('acao', 'Ação')}", expanded=False):
                        st.write(f"**Prioridade:** {rec.get('prioridade', 'N/A')}")
                        st.write(f"**Prazo:** {rec.get('prazo', 'N/A')}")
                        st.write(f"**Impacto:** {rec.get('impacto', 'N/A')}")
            
            # Outros insights estratégicos
            col1, col2 = st.columns(2)
            with col1:
                if estrategica.get("cenarios_futuros"):
                    st.subheader("🔮 Cenários Futuros")
                    for cenario in estrategica["cenarios_futuros"]:
                        st.markdown(f"• {cenario}")
            
            with col2:
                if estrategica.get("fatores_criticos_sucesso"):
                    st.subheader("🎯 Fatores Críticos de Sucesso")
                    for fator in estrategica["fatores_criticos_sucesso"]:
                        st.markdown(f"• {fator}")
            
            st.divider()
        
        # Síntese Final e Roadmap
        if analise.get("sintese"):
            sintese = analise["sintese"]
            
            # Conexões entre insights
            if sintese.get("conexoes_insights"):
                st.header("🔗 Conexões e Insights Integrados")
                for conexao in sintese["conexoes_insights"]:
                    st.info(f"💡 {conexao}")
            
            # Priorização de ações
            if sintese.get("priorizacao_acoes"):
                st.header("📋 Ações Priorizadas")
                acoes_ordenadas = sorted(sintese["priorizacao_acoes"], key=lambda x: x.get("score", 0), reverse=True)
                for acao in acoes_ordenadas:
                    score = acao.get("score", 0)
                    color = "🔴" if score >= 8 else "🟡" if score >= 5 else "🟢"
                    st.metric(
                        label=f"{color} {acao.get('acao', 'Ação')}",
                        value=f"Score: {score}/10",
                        help=acao.get('justificativa', '')
                    )
            
            # Roadmap
            if sintese.get("roadmap"):
                st.header("🗺️ Roadmap de Implementação")
                roadmap = sintese["roadmap"]
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    if roadmap.get("imediato"):
                        st.subheader("⚡ Imediato")
                        for acao in roadmap["imediato"]:
                            st.markdown(f"• {acao}")
                
                with col2:
                    if roadmap.get("30_dias"):
                        st.subheader("📅 30 Dias")
                        for acao in roadmap["30_dias"]:
                            st.markdown(f"• {acao}")
                
                with col3:
                    if roadmap.get("90_dias"):
                        st.subheader("🗓️ 90 Dias")
                        for acao in roadmap["90_dias"]:
                            st.markdown(f"• {acao}")
        
        # Debug info
        if st.session_state.dbg_cb_key:
            with st.expander("Debug: Análise Profunda Completa", expanded=False):
                st.json(analise)

if uploaded_file_sb is None and st.session_state.f_name is not None:
    keys_to_clear_on_remove = list(st.session_state.keys())
    preserved_widget_keys_on_remove = [
        "nav_radio_key_final_v7", # Atualize para as chaves únicas usadas
        "uploader_sidebar_key_final_v7", 
        "debug_cb_sidebar_key_final_v7"
    ] 
    if "s_gemini" in st.session_state: 
        for sug_key_cfg_clear in st.session_state.s_gemini:
            s_id_preserve_val_clear = sug_key_cfg_clear.get('id')
            if s_id_preserve_val_clear:
                preserved_widget_keys_on_remove.extend([f"acc_loop_{s_id_preserve_val_clear}", f"tit_loop_{s_id_preserve_val_clear}"])
            
    for key_cl_remove in keys_to_clear_on_remove:
        if key_cl_remove not in preserved_widget_keys_on_remove:
            if key_cl_remove in st.session_state: del st.session_state[key_cl_remove]
    
    for k_reinit_main, dv_reinit_main in [("s_gemini",[]),("cfg_sugs",{}),
                                ("doc_ctx",{"texto":"","tabelas":[]}),
                                ("f_name",None),("dbg_cb_key",False), 
                                ("pg_sel","Dashboard Principal"),("analise_profunda",{})]:
        st.session_state.setdefault(k_reinit_main, dv_reinit_main)
    st.rerun()
