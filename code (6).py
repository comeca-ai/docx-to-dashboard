import streamlit as st
from docx import Document
import pandas as pd
import plotly.express as px
import google.generativeai as genai
import json
import os
import traceback
import re 

# --- 1. Configuração da Chave da API do Gemini ---
def get_gemini_api_key():
    try: return st.secrets["GOOGLE_API_KEY"]
    except (FileNotFoundError, KeyError): 
        api_key = os.environ.get("GOOGLE_API_KEY")
        return api_key if api_key else None

# --- 2. Funções de Processamento do Documento e Interação com Gemini ---
def parse_value_for_numeric(val_str_in):
    if pd.isna(val_str_in) or str(val_str_in).strip() == '': return None
    text = str(val_str_in).strip()
    is_negative_paren = text.startswith('(') and text.endswith(')')
    if is_negative_paren: text = text[1:-1]
    text_num_part = re.sub(r'[R$\s%]', '', text)
    if ',' in text_num_part and '.' in text_num_part:
        if text_num_part.rfind('.') < text_num_part.rfind(','): text_num_part = text_num_part.replace('.', '') 
        text_num_part = text_num_part.replace(',', '.') 
    elif ',' in text_num_part: text_num_part = text_num_part.replace(',', '.')
    match = re.search(r"([-+]?\d*\.?\d+|\d+)", text_num_part)
    if match:
        try: num = float(match.group(1)); return -num if is_negative_paren else num
        except ValueError: return None
    return None

def extrair_conteudo_docx(uploaded_file):
    try:
        document = Document(uploaded_file)
        textos = [p.text for p in document.paragraphs if p.text.strip()]
        tabelas_data = [] 
        for i, table_obj in enumerate(document.tables):
            data_rows, keys, nome_tabela = [], None, f"Tabela_DOCX_{i+1}"
            try: 
                prev_el = table_obj._element.getprevious()
                if prev_el is not None and prev_el.tag.endswith('p'):
                    p_text = "".join(node.text for node in prev_el.xpath('.//w:t')).strip()
                    if p_text and len(p_text) < 80: nome_tabela = p_text.replace(":", "").strip()
            except: pass
            for r_idx, row in enumerate(table_obj.rows):
                cells = [c.text.strip() for c in row.cells]
                if r_idx == 0: keys = [k.replace("\n"," ").strip() if k else f"Col{c_idx+1}" for c_idx, k in enumerate(cells)]; continue
                if keys: 
                    row_dict = {}
                    for k_idx, key_name in enumerate(keys):
                        row_dict[key_name] = cells[k_idx] if k_idx < len(cells) else None
                    data_rows.append(row_dict)
            if data_rows:
                try:
                    df = pd.DataFrame(data_rows)
                    for col in df.columns:
                        original_series = df[col].copy()
                        num_series = original_series.astype(str).apply(parse_value_for_numeric)
                        if num_series.notna().sum() / max(1, len(num_series)) > 0.3:
                            df[col] = pd.to_numeric(num_series, errors='coerce')
                            continue 
                        else: df[col] = original_series 
                        try:
                            temp_str_col = df[col].astype(str)
                            # Removido infer_datetime_format=True pois é depreciado
                            dt_series = pd.to_datetime(temp_str_col, errors='coerce', dayfirst=True) 
                            if dt_series.notna().sum() / max(1, len(dt_series)) > 0.5:
                                df[col] = dt_series
                            else: 
                                df[col] = original_series.astype(str).fillna('')
                        except: df[col] = original_series.astype(str).fillna('')
                    for col in df.columns: 
                        if df[col].dtype == 'object': df[col] = df[col].astype(str).fillna('')
                    tabelas_data.append({"id": f"doc_tabela_{i+1}", "nome": nome_tabela, "dataframe": df})
                except Exception as e_df_proc:
                    st.warning(f"Não foi possível processar DataFrame para tabela '{nome_tabela}': {e_df_proc}")
        return "\n\n".join(textos), tabelas_data
    except Exception as e_doc_read: 
        st.error(f"Erro crítico ao ler arquivo DOCX: {e_doc_read}")
        return "", []

def analisar_documento_com_gemini(texto_doc, tabelas_info_list):
    api_key = get_gemini_api_key()
    if not api_key: st.warning("Chave API Gemini não configurada."); return []
    try:
        genai.configure(api_key=api_key)
        safety_settings = [{"category": c,"threshold": "BLOCK_NONE"} for c in ["HARM_CATEGORY_HARASSMENT","HARM_CATEGORY_HATE_SPEECH","HARM_CATEGORY_SEXUALLY_EXPLICIT","HARM_CATEGORY_DANGEROUS_CONTENT"]]
        model = genai.GenerativeModel(model_name="gemini-1.5-flash-latest", safety_settings=safety_settings)
        
        tabelas_prompt_str = ""
        for t_info in tabelas_info_list:
            df, nome_t, id_t = t_info["dataframe"], t_info["nome"], t_info["id"]
            sample_df = df.head(3).iloc[:, :min(5, len(df.columns))]
            md_table = ""
            try: md_table = sample_df.to_markdown(index=False)
            except: md_table = sample_df.to_string(index=False) 
            
            colunas_para_mostrar_tipos = df.columns.tolist()[:min(8, len(df.columns))]
            col_types_list = [f"'{col_name_prompt}' (tipo: {str(df[col_name_prompt].dtype)})" for col_name_prompt in colunas_para_mostrar_tipos]
            col_types_str = ", ".join(col_types_list)
            
            tabelas_prompt_str += f"\n--- Tabela '{nome_t}' (ID: {id_t}) ---\nColunas e tipos (primeiras {len(colunas_para_mostrar_tipos)}): {col_types_str}\nAmostra de dados:\n{md_table}\n"
        
        text_limit = 45000
        prompt_text = texto_doc[:text_limit] + ("\n[TEXTO TRUNCADO...]" if len(texto_doc) > text_limit else "")
        
        prompt = f"""
        Você é um assistente de análise de dados. Analise o texto e as tabelas.
        [TEXTO]{prompt_text}[FIM TEXTO]
        [TABELAS]{tabelas_prompt_str}[FIM TABELAS]

        Gere lista JSON de sugestões de visualizações. Objeto DEVE ter: "id", "titulo", "tipo_sugerido" ("kpi", "tabela_dados", "lista_swot", "grafico_barras", "grafico_pizza", "grafico_linha", "grafico_dispersao", "grafico_barras_agrupadas"), "fonte_id" (ID tabela ou "texto_descricao_fonte"), "parametros" (objeto JSON), "justificativa".
        Para "parametros":
        - "kpi": {{"valor": "ValorKPI", "delta": "Mudança", "descricao": "Contexto"}}
        - "tabela_dados": Para TABELA EXISTENTE: {{"id_tabela_original": "ID_Tabela"}}. Para DADOS DO TEXTO: {{"dados": [{{"Coluna1": "ValorA1"}}, ...], "colunas_titulo": ["Título Col1"]}}
        - "lista_swot": {{"forcas": ["F1"], "fraquezas": ["Fr1"], "oportunidades": ["Op1"], "ameacas": ["Am1"]}} (Listas de strings).
        - Gráficos de TABELA ("barras", "linha", "dispersao"): {{"eixo_x": "NOME_COL_X", "eixo_y": "NOME_COL_Y"}} (Y numérico, use nomes exatos).
        - Gráficos de PIZZA de TABELA: {{"categorias": "NOME_COL_CAT", "valores": "NOME_COL_VAL_NUM"}} (Valores numéricos, use nomes exatos).
        - Gráficos com DADOS EXTRAÍDOS DO TEXTO ("barras", "pizza", etc.): {{"dados": [{{"NomeEixoX": "CatA", "NomeEixoY": ValNumA}}, ...], "eixo_x": "NomeEixoX", "eixo_y": "NomeEixoY"}} (Valores DEVEM ser numéricos).
        - "grafico_barras_agrupadas": Se de TABELA: {{"eixo_x": "COL_PRINCIPAL", "eixo_y": "COL_VALOR_NUM", "cor_agrupamento": "COL_SUB_CAT"}}. Se DADOS EXTRAÍDOS: {{"dados": [{{"CatPrincipal": "A", "SubCat": "X", "Valor": 10}}, ...], "eixo_x": "CatPrincipal", "eixo_y": "Valor", "cor_agrupamento": "SubCat"}}.
        
        INSTRUÇÕES CRÍTICAS:
        1.  NOMES DE COLUNAS: Para gráficos de TABELA, use os NOMES EXATOS das colunas como fornecidos nos "Colunas e tipos".
        2.  DADOS NUMÉRICOS: Se a coluna de valor de uma TABELA não for numérica (float64/int64) conforme os "tipos inferidos", NÃO sugira gráfico que exija valor numérico para ela, A MENOS que você possa confiavelmente extrair um valor numérico do seu conteúdo textual (ex: de '70%' extrair 70.0; de '70% - 86%' extrair 70.0 ou 78.0; de 'R$ 15,5 Bi' extrair 15.5). Se extrair do texto, coloque em "dados" e certifique-se que os valores sejam números, não strings de números.
        3.  COBERTURA GEOGRÁFICA (Player, Cidades): Se for apenas lista, sugira "tabela_dados" e forneça os dados extraídos no campo "dados" dos "parametros" com "colunas_titulo".
        4.  SWOT: Se uma tabela compara SWOTs (ex: Tabela 4 do documento), gere sugestões "lista_swot" INDIVIDUAL para CADA player listado nessa tabela, usando o nome do player no "titulo". Se o SWOT estiver no texto, extraia os pontos para "forcas", "fraquezas", etc.
        Retorne APENAS a lista JSON válida. Seja conciso na justificativa.
        """
        with st.spinner("🤖 Gemini analisando... (Pode levar alguns instantes)"):
            response = model.generate_content(prompt)
        cleaned_text = response.text.strip().lstrip("```json").rstrip("```").strip()
        sugestoes = json.loads(cleaned_text)
        if isinstance(sugestoes, list) and all(isinstance(item, dict) for item in sugestoes):
             st.success(f"{len(sugestoes)} sugestões!"); return sugestoes
        st.error("Resposta Gemini não é lista JSON."); return []
    except json.JSONDecodeError as e: st.error(f"Erro JSON Gemini: {e}"); st.code(response.text if 'response' in locals() else "N/A", language="text"); return []
    except Exception as e: st.error(f"Erro API Gemini: {e}"); st.text(traceback.format_exc()); return []

def render_kpis(kpi_sugestoes):
    if kpi_sugestoes:
        num_kpis = len(kpi_sugestoes); kpi_cols = st.columns(min(num_kpis, 4)) 
        for i, kpi_sug in enumerate(kpi_sugestoes):
            with kpi_cols[i % min(num_kpis, 4)]:
                params=kpi_sug.get("parametros",{}); delta_val=str(params.get("delta",""))
                st.metric(label=kpi_sug.get("titulo","KPI"),value=str(params.get("valor","N/A")),delta=delta_val if delta_val else None,help=params.get("descricao"))
        st.divider()

def render_swot_card(titulo_completo_swot, swot_data, card_key_prefix=""):
    st.subheader(f"{titulo_completo_swot}") 
    col1, col2 = st.columns(2)
    swot_map = {"forcas": ("Forças 💪", col1), "fraquezas": ("Fraquezas 📉", col1), 
                "oportunidades": ("Oportunidades 🚀", col2), "ameacas": ("Ameaças ⚠️", col2)}
    for key_swot_category, (header_swot_render, col_target_swot_render) in swot_map.items():
        with col_target_swot_render:
            st.markdown(f"##### {header_swot_render}")
            points_swot_render = swot_data.get(key_swot_category, ["N/A (info. não fornecida)"])
            if not points_swot_render or not isinstance(points_swot_render, list) or not all(isinstance(p_swot, str) for p_swot in points_swot_render): 
                points_swot_render = ["N/A (formato de dados incorreto)"]
            if not points_swot_render: points_swot_render = ["N/A"] 
            for item_swot_render in points_swot_render: 
                st.markdown(f"<div style='margin-bottom: 5px;'>- {item_swot_render}</div>", unsafe_allow_html=True)
    st.markdown("<hr style='margin-top: 10px; margin-bottom: 20px;'>", unsafe_allow_html=True)

def render_plotly_chart(item_config, df_plot_input):
    if df_plot_input is None:
        st.warning(f"Dados não disponíveis para o gráfico '{item_config.get('titulo', 'Sem Título')}'.")
        return False
    df_plot = df_plot_input.copy()
    tipo_grafico, titulo, params = item_config.get("tipo_sugerido"), item_config.get("titulo"), item_config.get("parametros", {})
    x_col, y_col, cat_col, val_col = params.get("eixo_x"), params.get("eixo_y"), params.get("categorias"), params.get("valores")
    cor_agrupamento_col = params.get("cor_agrupamento")
    fig, plot_func, plot_args = None, None, {}
    if tipo_grafico in ["grafico_barras", "grafico_barras_agrupadas"] and x_col and y_col: 
        plot_func,plot_args=px.bar,{"x":x_col,"y":y_col}
        if tipo_grafico == "grafico_barras_agrupadas" and cor_agrupamento_col:
            plot_args["color"], plot_args["barmode"] = cor_agrupamento_col, "group"
    elif tipo_grafico=="grafico_linha" and x_col and y_col: plot_func,plot_args=px.line,{"x":x_col,"y":y_col,"markers":True}
    elif tipo_grafico=="grafico_dispersao" and x_col and y_col: plot_func,plot_args=px.scatter,{"x":x_col,"y":y_col}
    elif tipo_grafico=="grafico_pizza" and cat_col and val_col: plot_func,plot_args=px.pie,{"names":cat_col,"values":val_col}
    if plot_func:
        required_cols=[col for col in plot_args.values() if isinstance(col,str)]
        if not all(col in df_plot.columns for col in required_cols):
            st.warning(f"Colunas {required_cols} não encontradas para '{titulo}'. Disponíveis: {df_plot.columns.tolist()}")
            return False
        try:
            df_plot_cleaned = df_plot.copy() 
            y_ax,val_ax=plot_args.get("y"),plot_args.get("values")
            if y_ax and y_ax in df_plot_cleaned.columns: df_plot_cleaned[y_ax]=pd.to_numeric(df_plot_cleaned[y_ax],errors='coerce')
            if val_ax and val_ax in df_plot_cleaned.columns: df_plot_cleaned[val_ax]=pd.to_numeric(df_plot_cleaned[val_ax],errors='coerce')
            cols_to_check_na_final = [val_check for val_check in plot_args.values() if isinstance(val_check, str) and val_check in df_plot_cleaned.columns]
            df_plot_cleaned.dropna(subset=cols_to_check_na_final, inplace=True)
            if not df_plot_cleaned.empty:
                fig=plot_func(df_plot_cleaned,title=titulo,**plot_args); st.plotly_chart(fig,use_container_width=True); return True
            else: st.warning(f"Dados insuficientes para '{titulo}' após limpar NaNs de {cols_to_check_na_final}.")
        except Exception as e_plotly_render: 
            st.warning(f"Erro ao gerar gráfico Plotly '{titulo}': {e_plotly_render}. Dtypes: {df_plot.dtypes.to_dict() if df_plot is not None else 'DF é None'}")
    elif tipo_grafico in ["grafico_barras","grafico_barras_agrupadas","grafico_linha","grafico_dispersao","grafico_pizza","grafico_radar"]:
        st.warning(f"Configuração de parâmetros incompleta para '{titulo}' (tipo: {tipo_grafico}).")
    return False

# --- 3. Interface Streamlit Principal ---
st.set_page_config(layout="wide", page_title="Gemini DOCX Insights")
for k, dv in [("sugestoes_gemini",[]),("config_sugestoes",{}),("conteudo_docx",{"texto":"","tabelas":[]}),
              ("nome_arquivo_atual",None),("debug_checkbox_key",False),("pagina_selecionada","Dashboard Principal")]:
    st.session_state.setdefault(k, dv)

st.sidebar.title("✨ Navegação"); pagina_opcoes_sidebar = ["Dashboard Principal", "Análise SWOT Detalhada"]
st.session_state.pagina_selecionada = st.sidebar.radio(
    "Selecione:", pagina_opcoes_sidebar, 
    index=pagina_opcoes_sidebar.index(st.session_state.pagina_selecionada), 
    key="nav_radio_key_final_v3" 
)
st.sidebar.divider(); uploaded_file_sidebar = st.sidebar.file_uploader("Selecione DOCX", type="docx", key="uploader_sidebar_key_final_v3")
show_debug_info_sidebar = st.sidebar.checkbox("Mostrar Informações de Depuração", 
                                    value=st.session_state.debug_checkbox_key, 
                                    key="debug_cb_sidebar_key_final_v3") 
st.session_state.debug_checkbox_key = show_debug_info_sidebar

if uploaded_file_sidebar:
    if st.session_state.nome_arquivo_atual != uploaded_file_sidebar.name: 
        with st.spinner("Processando novo documento..."):
            st.session_state.sugestoes_gemini, st.session_state.config_sugestoes = [], {}
            st.session_state.nome_arquivo_atual = uploaded_file_sidebar.name
            texto_doc_main, tabelas_doc_main = extrair_conteudo_docx(uploaded_file_sidebar)
            st.session_state.conteudo_docx = {"texto": texto_doc_main, "tabelas": tabelas_doc_main}
            if texto_doc_main or tabelas_doc_main:
                sugestoes_main = analisar_documento_com_gemini(texto_doc_main, tabelas_doc_main)
                st.session_state.sugestoes_gemini = sugestoes_main
                temp_config_init_main = {}
                for i_init_main,s_init_main in enumerate(sugestoes_main): 
                    s_id_init_main = s_init_main.get("id", f"s_init_main_{i_init_main}_{hash(s_init_main.get('titulo',''))}"); s_init_main["id"] = s_id_init_main
                    temp_config_init_main[s_id_init_main] = {"aceito":True,"titulo_editado":s_init_main.get("titulo","S/Título"),"dados_originais":s_init_main}
                st.session_state.config_sugestoes = temp_config_init_main
            else: st.sidebar.warning("Nenhum conteúdo extraído do DOCX.")
    
    if show_debug_info_sidebar and (st.session_state.conteudo_docx["texto"] or st.session_state.conteudo_docx["tabelas"]):
        with st.expander("Debug: Conteúdo DOCX (após extração e tipos)", expanded=False):
            st.text_area("Texto (amostra)", st.session_state.conteudo_docx["texto"][:1000], height=80)
            for t_info_dbg_main in st.session_state.conteudo_docx["tabelas"]:
                st.write(f"ID: {t_info_dbg_main['id']}, Nome: {t_info_dbg_main['nome']}")
                try: st.dataframe(t_info_dbg_main['dataframe'].head().astype(str).fillna("-")) 
                except Exception: st.text(f"Head:\n{t_info_dbg_main['dataframe'].head().to_string(na_rep='-')}")
                st.write("Tipos:", t_info_dbg_main['dataframe'].dtypes.to_dict())

    if st.session_state.sugestoes_gemini:
        st.sidebar.divider(); st.sidebar.header("⚙️ Configurar Sugestões")
        for sug_cfg_sidebar in st.session_state.sugestoes_gemini:
            s_id_cfg_sb = sug_cfg_sidebar.get('id') 
            if not s_id_cfg_sb : continue 
            if s_id_cfg_sb not in st.session_state.config_sugestoes:
                 st.session_state.config_sugestoes[s_id_cfg_sb] = {"aceito":True,"titulo_editado":sug_cfg_sidebar.get("titulo","S/Título"),"dados_originais":sug_cfg_sidebar}
            cfg_current_sb = st.session_state.config_sugestoes[s_id_cfg_sb]
            
            with st.sidebar.expander(f"{cfg_current_sb['titulo_editado']}",expanded=False):
                st.caption(f"Tipo: {sug_cfg_sidebar.get('tipo_sugerido')} | Fonte: {sug_cfg_sidebar.get('fonte_id')}")
                cfg_current_sb["aceito"]=st.checkbox("Incluir?",value=cfg_current_sb["aceito"],key=f"acc_cfg_{s_id_cfg_sb}")
                cfg_current_sb["titulo_editado"]=st.text_input("Título",value=cfg_current_sb["titulo_editado"],key=f"tit_cfg_{s_id_cfg_sb}")
else: 
    if st.session_state.pagina_selecionada == "Dashboard Principal":
        st.info("Por favor, faça o upload de um arquivo DOCX na barra lateral para começar.")

# --- RENDERIZAÇÃO DA PÁGINA SELECIONADA ---
if st.session_state.pagina_selecionada == "Dashboard Principal":
    st.title("📊 Dashboard de Insights do Documento")
    if uploaded_file_sidebar and st.session_state.sugestoes_gemini:
        kpis_render, outros_render = [], []
        for s_id_main_dash, s_cfg_main_dash in st.session_state.config_sugestoes.items():
            if s_cfg_main_dash["aceito"]: 
                item_main_dash = {"titulo":s_cfg_main_dash["titulo_editado"], **s_cfg_main_dash["dados_originais"]}
                (kpis_render if item_main_dash.get("tipo_sugerido")=="kpi" else outros_render).append(item_main_dash)
        
        render_kpis(kpis_render)
        
        if show_debug_info_sidebar:
             with st.expander("Debug: Elementos para Dashboard Principal (Não-KPI)", expanded=True): # Expandido por padrão
                st.json({"Outros Elementos (Configurados e Aceitos)": outros_render}, expanded=False)
        
        elementos_renderizados_dash = 0 
        col_idx_dash = 0 
        if outros_render:
            item_cols_main_dash = st.columns(2)
            for item_render_loop in outros_render:
                if item_render_loop.get("tipo_sugerido") == "lista_swot": continue 
                
                with item_cols_main_dash[col_idx_dash % 2]:
                    df_plot_loop, rendered_loop = None, False
                    params_loop = item_render_loop.get("parametros",{})
                    tipo_loop = item_render_loop.get("tipo_sugerido")
                    fonte_loop = item_render_loop.get("fonte_id")
                    titulo_loop = item_render_loop.get("titulo", "Visualização")
                    
                    st.subheader(titulo_loop) 
                    try:
                        if params_loop.get("dados"): 
                            try: df_plot_loop=pd.DataFrame(params_loop["dados"])
                            except Exception as e_dfd_loop: st.warning(f"'{titulo_loop}': Erro DF 'dados': {e_dfd_loop}"); continue
                        elif str(fonte_loop).startswith("doc_tabela_"): 
                            df_plot_loop=next((t["dataframe"] for t in st.session_state.conteudo_docx["tabelas"] if t["id"]==fonte_loop),None)
                        
                        if tipo_loop=="tabela_dados":
                            df_tbl_loop=None
                            if str(fonte_loop).startswith("texto_") and params_loop.get("dados"):
                                try: 
                                    df_tbl_loop=pd.DataFrame(params_loop.get("dados")); 
                                    if params_loop.get("colunas_titulo"): df_tbl_loop.columns=params_loop.get("colunas_titulo")
                                except Exception as e_dftxt_loop: st.warning(f"Erro tabela texto '{titulo_loop}': {e_dftxt_loop}")
                            else: 
                                id_tbl_loop=params_loop.get("id_tabela_original",fonte_loop)
                                df_tbl_loop=next((t["dataframe"] for t in st.session_state.conteudo_docx["tabelas"] if t["id"]==id_tbl_loop),None)
                            
                            if df_tbl_loop is not None: 
                                try: st.dataframe(df_tbl_loop.astype(str).fillna("-"))
                                except: st.text(df_tbl_loop.to_string(na_rep='-')); 
                                rendered_loop=True
                            else: st.warning(f"Tabela '{titulo_loop}' (Fonte: {fonte_loop}) não encontrada.")
                        
                        elif tipo_loop in ["grafico_barras","grafico_linha","grafico_dispersao","grafico_pizza", "grafico_barras_agrupadas"]:
                            if render_plotly_chart(item_render_loop, df_plot_loop): rendered_loop = True
                        
                        elif tipo_loop == 'mapa': 
                            st.info(f"Mapa para '{titulo_loop}' não implementado.")
                            rendered_loop=True
                        
                        if not rendered_loop and tipo_loop not in ["kpi","lista_swot","mapa"]: 
                            st.info(f"'{titulo_loop}' ({tipo_loop}) não gerado. Dados/Tipo não suportado.")
                    except Exception as e_render_loop: 
                        st.error(f"Erro renderizando '{titulo_loop}': {e_render_loop}")
                
                if rendered_loop: 
                    idx_main_dash+=1
                    elementos_renderizados_dash+=1
            
            if elementos_renderizados_dash == 0 and any(c['aceito'] and c['dados_originais'].get('tipo_sugerido') not in ['kpi','lista_swot'] for c in st.session_state.config_sugestoes.values()):
                st.info("Nenhum gráfico/tabela (além de KPIs/SWOTs) pôde ser gerado para o Dashboard Principal.")
        
        elif not kpis_render and not uploaded_file_sidebar: pass 
        elif not kpis_render and not outros_render and uploaded_file_sidebar and st.session_state.sugestoes_gemini: 
            st.info("Nenhum elemento selecionado ou gerado para o dashboard principal.")

elif st.session_state.pagina_selecionada == "Análise SWOT Detalhada":
    st.title("🔬 Análise SWOT Detalhada")
    if not uploaded_file_sidebar: st.warning("Faça upload de um DOCX na barra lateral.")
    elif not st.session_state.sugestoes_gemini: st.info("Aguardando processamento ou nenhuma sugestão gerada.")
    else:
        swot_sugs_page = [s_cfg_swot_page["dados_originais"] for s_id_swot_page,s_cfg_swot_page in st.session_state.config_sugestoes.items() 
                        if s_cfg_swot_page["aceito"] and s_cfg_swot_page["dados_originais"].get("tipo_sugerido")=="lista_swot"]
        if not swot_sugs_page: st.info("Nenhuma análise SWOT sugerida/selecionada para esta página.")
        else:
            if show_debug_info_sidebar:
                with st.expander("Debug: Dados para Análise SWOT (Página Dedicada)", expanded=False):
                    st.json({"SWOTs Selecionados para esta página": swot_sugs_page})
            for swot_item_render_page in swot_sugs_page:
                render_swot_card(
                    swot_item_render_page.get("titulo","Análise SWOT"), 
                    swot_item_render_page.get("parametros",{}), 
                    card_key_prefix=swot_item_render_page.get("id","swot_page_default") 
                )

if uploaded_file_sidebar is None and st.session_state.nome_arquivo_atual is not None:
    keys_to_clear_on_remove = list(st.session_state.keys())
    preserved_widget_keys_on_remove = [
        "nav_radio_key_final_v3", "uploader_sidebar_key_final_v3", "debug_cb_sidebar_key_final_v3"
    ] 
    for sug_key_cfg_clear in st.session_state.get("sugestoes_gemini", []): # Itera sobre cópia ou usa .get
        s_id_preserve_val_clear = sug_key_cfg_clear.get('id')
        if s_id_preserve_val_clear:
            preserved_widget_keys_on_remove.extend([f"acc_cfg_{s_id_preserve_val_clear}", f"tit_cfg_{s_id_preserve_val_clear}"])
            
    for key_cl_remove in keys_to_clear_on_remove:
        if key_cl_remove not in preserved_widget_keys_on_remove:
            if key_cl_remove in st.session_state: del st.session_state[key_cl_remove]
    
    for k_reinit_remove, dv_reinit_remove in [("sugestoes_gemini",[]),("config_sugestoes",{}),
                                ("conteudo_docx",{"texto":"","tabelas":[]}),
                                ("nome_arquivo_atual",None),("debug_checkbox_key",False), 
                                ("pagina_selecionada","Dashboard Principal")]:
        st.session_state.setdefault(k_reinit_remove, dv_reinit_remove)
    st.experimental_rerun()