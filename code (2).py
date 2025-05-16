import streamlit as st
from docx import Document
import pandas as pd
import plotly.express as px
import google.generativeai as genai
import json
import os
import traceback
import re 

# --- 1. Configuração da Chave da API do Gemini ---
def get_gemini_api_key():
    try: return st.secrets["GOOGLE_API_KEY"]
    except (FileNotFoundError, KeyError): 
        api_key = os.environ.get("GOOGLE_API_KEY")
        return api_key if api_key else None

# --- 2. Funções de Processamento do Documento e Interação com Gemini ---
def parse_value_for_numeric(val_str_in):
    if pd.isna(val_str_in) or str(val_str_in).strip() == '': return None
    text = str(val_str_in).strip()
    is_negative_paren = text.startswith('(') and text.endswith(')')
    if is_negative_paren: text = text[1:-1]
    text_num_part = re.sub(r'[R$\s%]', '', text)
    if ',' in text_num_part and '.' in text_num_part:
        if text_num_part.rfind('.') < text_num_part.rfind(','): text_num_part = text_num_part.replace('.', '') 
        text_num_part = text_num_part.replace(',', '.') 
    elif ',' in text_num_part: text_num_part = text_num_part.replace(',', '.')
    match = re.search(r"([-+]?\d*\.?\d+|\d+)", text_num_part)
    if match:
        try: num = float(match.group(1)); return -num if is_negative_paren else num
        except ValueError: return None
    return None

def extrair_conteudo_docx(uploaded_file):
    try:
        document = Document(uploaded_file)
        textos = [p.text for p in document.paragraphs if p.text.strip()]
        tabelas_data = [] 
        for i, table_obj in enumerate(document.tables):
            data_rows, keys, nome_tabela = [], None, f"Tabela_DOCX_{i+1}"
            try: 
                prev_el = table_obj._element.getprevious()
                if prev_el is not None and prev_el.tag.endswith('p'):
                    p_text = "".join(node.text for node in prev_el.xpath('.//w:t')).strip()
                    if p_text and len(p_text) < 80: nome_tabela = p_text.replace(":", "").strip()
            except: pass
            for r_idx, row in enumerate(table_obj.rows):
                cells = [c.text.strip() for c in row.cells]
                if r_idx == 0: keys = [k.replace("\n"," ").strip() if k else f"Col{c_idx+1}" for c_idx, k in enumerate(cells)]; continue
                if keys: 
                    row_dict = {}
                    for k_idx, key_name in enumerate(keys):
                        row_dict[key_name] = cells[k_idx] if k_idx < len(cells) else None
                    data_rows.append(row_dict)
            if data_rows:
                df = pd.DataFrame(data_rows)
                for col in df.columns:
                    original_series = df[col].copy()
                    num_series = original_series.astype(str).apply(parse_value_for_numeric)
                    if num_series.notna().sum() / max(1, len(num_series)) > 0.3:
                        df[col] = pd.to_numeric(num_series, errors='coerce')
                        continue 
                    else: df[col] = original_series 
                    try:
                        temp_str_col = df[col].astype(str)
                        dt_series = pd.to_datetime(temp_str_col, errors='coerce', dayfirst=True) 
                        if dt_series.notna().sum() / max(1, len(dt_series)) > 0.5:
                            df[col] = dt_series
                        else: 
                            df[col] = original_series.astype(str).fillna('')
                    except: df[col] = original_series.astype(str).fillna('')
                for col in df.columns: 
                    if df[col].dtype == 'object': df[col] = df[col].astype(str).fillna('')
                tabelas_data.append({"id": f"doc_tabela_{i+1}", "nome": nome_tabela, "dataframe": df})
        return "\n\n".join(textos), tabelas_data
    except Exception as e: st.error(f"Erro crítico ao ler DOCX: {e}"); return "", []

def analisar_documento_com_gemini(texto_doc, tabelas_info_list):
    api_key = get_gemini_api_key()
    if not api_key: st.warning("Chave API Gemini não configurada."); return []
    try:
        genai.configure(api_key=api_key)
        safety_settings = [{"category": c,"threshold": "BLOCK_NONE"} for c in ["HARM_CATEGORY_HARASSMENT","HARM_CATEGORY_HATE_SPEECH","HARM_CATEGORY_SEXUALLY_EXPLICIT","HARM_CATEGORY_DANGEROUS_CONTENT"]]
        model = genai.GenerativeModel(model_name="gemini-1.5-flash-latest", safety_settings=safety_settings)
        
        tabelas_prompt_str = ""
        for t_info in tabelas_info_list:
            df, nome_t, id_t = t_info["dataframe"], t_info["nome"], t_info["id"]
            sample_df = df.head(3).iloc[:, :min(5, len(df.columns))]
            md_table = ""
            try: md_table = sample_df.to_markdown(index=False)
            except: md_table = sample_df.to_string(index=False) # Fallback
            
            colunas_para_mostrar_tipos = df.columns.tolist()[:min(8, len(df.columns))]
            col_types_list = [f"'{col_name_prompt}' (tipo: {str(df[col_name_prompt].dtype)})" for col_name_prompt in colunas_para_mostrar_tipos]
            col_types_str = ", ".join(col_types_list)
            
            tabelas_prompt_str += f"\n--- Tabela '{nome_t}' (ID: {id_t}) ---\nColunas e tipos (primeiras {len(colunas_para_mostrar_tipos)}): {col_types_str}\nAmostra de dados:\n{md_table}\n"
        
        text_limit = 45000
        prompt_text = texto_doc[:text_limit] + ("\n[TEXTO TRUNCADO...]" if len(texto_doc) > text_limit else "")
        
        prompt = f"""
        Você é um assistente de análise de dados. Analise o texto e as tabelas.
        [TEXTO]{prompt_text}[FIM TEXTO]
        [TABELAS]{tabelas_prompt_str}[FIM TABELAS]

        Gere lista JSON de sugestões de visualizações. Objeto DEVE ter: "id", "titulo", "tipo_sugerido" ("kpi", "tabela_dados", "lista_swot", "grafico_barras", "grafico_pizza", "grafico_linha", "grafico_dispersao"), "fonte_id" (ID tabela ou "texto_descricao_fonte"), "parametros" (objeto JSON), "justificativa".
        Para "parametros":
        - "kpi": {{"valor": "ValorKPI", "delta": "Mudança", "descricao": "Contexto"}}
        - "tabela_dados": Para TABELA EXISTENTE: {{"id_tabela_original": "ID_Tabela"}}. Para DADOS DO TEXTO: {{"dados": [{{"Coluna1": "ValorA1"}}, ...], "colunas_titulo": ["Título Col1"]}}
        - "lista_swot": {{"forcas": ["F1"], "fraquezas": ["Fr1"], "oportunidades": ["Op1"], "ameacas": ["Am1"]}} (Listas de strings).
        - Gráficos de TABELA ("barras", "linha", "dispersao"): {{"eixo_x": "NOME_COL_X", "eixo_y": "NOME_COL_Y"}} (Y numérico, use nomes exatos).
        - Gráficos de PIZZA de TABELA: {{"categorias": "NOME_COL_CAT", "valores": "NOME_COL_VAL_NUM"}} (Valores numéricos, use nomes exatos).
        - Gráficos com DADOS EXTRAÍDOS DO TEXTO: {{"dados": [{{"NomeEixoX": "CatA", "NomeEixoY": ValNumA}}, ...], "eixo_x": "NomeEixoX", "eixo_y": "NomeEixoY"}} (Valores DEVEM ser numéricos).
        
        INSTRUÇÕES CRÍTICAS:
        1.  NOMES DE COLUNAS: Para gráficos de TABELA, use os NOMES EXATOS das colunas como fornecidos nos "Colunas e tipos".
        2.  DADOS NUMÉRICOS: Se a coluna de valor de uma TABELA não for numérica (float64/int64) conforme os "tipos inferidos", NÃO sugira gráfico que exija valor numérico para ela, A MENOS que você possa confiavelmente extrair um valor numérico do seu conteúdo textual (ex: de '70%' extrair 70.0; de '70% - 86%' extrair 70.0 ou 78.0; de 'R$ 15,5 Bi' extrair 15.5). Se extrair do texto, coloque em "dados" e certifique-se que os valores sejam números, não strings de números.
        3.  COBERTURA GEOGRÁFICA (Player, Cidades): Se for apenas lista, sugira "tabela_dados" e forneça os dados extraídos no campo "dados" dos "parametros" com "colunas_titulo".
        4.  SWOT: Se uma tabela compara SWOTs (ex: Tabela 4 do documento), gere sugestões "lista_swot" INDIVIDUAL para CADA player listado nessa tabela, usando o nome do player no "titulo". Se o SWOT estiver no texto, extraia os pontos para "forcas", "fraquezas", etc.
        Retorne APENAS a lista JSON válida. Seja conciso na justificativa.
        """
        with st.spinner("🤖 Gemini analisando... (Pode levar alguns instantes)"):
            # st.text_area("Debug Prompt:", prompt, height=150) 
            response = model.generate_content(prompt)
        cleaned_text = response.text.strip().lstrip("```json").rstrip("```").strip()
        # st.text_area("Debug Resposta Gemini:", cleaned_text, height=150)
        sugestoes = json.loads(cleaned_text)
        if isinstance(sugestoes, list) and all(isinstance(item, dict) for item in sugestoes):
             st.success(f"{len(sugestoes)} sugestões!"); return sugestoes
        st.error("Resposta Gemini não é lista JSON."); return []
    except json.JSONDecodeError as e: st.error(f"Erro JSON Gemini: {e}"); st.code(response.text if 'response' in locals() else "N/A", language="text"); return []
    except Exception as e: st.error(f"Erro API Gemini: {e}"); st.text(traceback.format_exc()); return []

# --- Funções de Renderização Específicas ---
def render_kpis(kpi_sugestoes):
    if kpi_sugestoes:
        num_kpis = len(kpi_sugestoes)
        kpi_cols = st.columns(min(num_kpis, 4)) 
        for i, kpi_sug in enumerate(kpi_sugestoes):
            with kpi_cols[i % min(num_kpis, 4)]:
                params = kpi_sug.get("parametros",{})
                st.metric(label=kpi_sug.get("titulo","KPI"), value=str(params.get("valor", "N/A")),
                          delta=str(params.get("delta", "")), help=params.get("descricao"))
        st.divider()

def render_swot_card(titulo_completo_swot, swot_data, card_key_prefix=""):
    st.subheader(f"{titulo_completo_swot}") 
    col1, col2 = st.columns(2)
    swot_map = {"forcas": ("Forças 💪", col1), "fraquezas": ("Fraquezas 📉", col1), 
                "oportunidades": ("Oportunidades 🚀", col2), "ameacas": ("Ameaças ⚠️", col2)}
    for key, (header, col_target) in swot_map.items():
        with col_target:
            st.markdown(f"##### {header}")
            points = swot_data.get(key, ["N/A (info. não fornecida)"])
            if not points or not isinstance(points, list) or not all(isinstance(p, str) for p in points): 
                points = ["N/A (formato de dados incorreto)"]
            if not points: points = ["N/A"] # Garante que não é lista vazia
            for point_idx, item_swot in enumerate(points): 
                st.markdown(f"<div style='margin-bottom: 5px;'>- {item_swot}</div>", unsafe_allow_html=True, 
                            key=f"swot_item_{card_key_prefix}_{titulo_completo_swot.replace(' ','_')}_{key}_{point_idx}")
    st.markdown("<hr style='margin-top: 10px; margin-bottom: 20px;'>", unsafe_allow_html=True)

def render_plotly_chart(item_config, df_plot_input):
    if df_plot_input is None:
        st.warning(f"Dados não disponíveis para o gráfico '{item_config.get('titulo', 'Sem Título')}'.")
        return False
    df_plot = df_plot_input.copy()
    tipo_grafico, titulo, params = item_config.get("tipo_sugerido"), item_config.get("titulo"), item_config.get("parametros", {})
    x_col, y_col, cat_col, val_col = params.get("eixo_x"), params.get("eixo_y"), params.get("categorias"), params.get("valores")
    fig, plot_func, plot_args = None, None, {}
    if tipo_grafico=="grafico_barras" and x_col and y_col: plot_func,plot_args=px.bar,{"x":x_col,"y":y_col}
    elif tipo_grafico=="grafico_linha" and x_col and y_col: plot_func,plot_args=px.line,{"x":x_col,"y":y_col,"markers":True}
    elif tipo_grafico=="grafico_dispersao" and x_col and y_col: plot_func,plot_args=px.scatter,{"x":x_col,"y":y_col}
    elif tipo_grafico=="grafico_pizza" and cat_col and val_col: plot_func,plot_args=px.pie,{"names":cat_col,"values":val_col}
    if plot_func:
        required_cols=[col for col in plot_args.values() if isinstance(col,str)]
        if not all(col in df_plot.columns for col in required_cols):
            st.warning(f"Colunas {required_cols} não encontradas para '{titulo}'. Disponíveis: {df_plot.columns.tolist()}")
            return False
        try:
            y_ax,val_ax=plot_args.get("y"),plot_args.get("values")
            if y_ax and y_ax in df_plot.columns: df_plot[y_ax]=pd.to_numeric(df_plot[y_ax],errors='coerce')
            if val_ax and val_ax in df_plot.columns: df_plot[val_ax]=pd.to_numeric(df_plot[val_ax],errors='coerce')
            df_plot_cleaned = df_plot.dropna(subset=required_cols).copy()
            if not df_plot_cleaned.empty:
                fig=plot_func(df_plot_cleaned,title=titulo,**plot_args); st.plotly_chart(fig,use_container_width=True); return True
            else: st.warning(f"Dados insuficientes para '{titulo}' após limpar NaNs de {required_cols}.")
        except Exception as e: st.warning(f"Erro Plotly '{titulo}': {e}. Dtypes: {df_plot.dtypes.to_dict()}")
    elif tipo_grafico in ["grafico_barras","grafico_linha","grafico_dispersao","grafico_pizza"]:
        st.warning(f"Params incompletos para '{titulo}' ({tipo_grafico}).")
    return False

# --- 3. Interface Streamlit Principal ---
st.set_page_config(layout="wide", page_title="Gemini DOCX Insights")
for k,dv in [("s_gemini",[]),("cfg_sugs",{}),("doc_ctx",{"texto":"","tabelas":[]}),("f_name",None),("dbg_cb_key",False),("pg_sel","Dashboard Principal")]: st.session_state.setdefault(k,dv)

st.sidebar.title("✨ Navegação"); pg_opts=["Dashboard Principal","Análise SWOT Detalhada"]
st.session_state.pg_sel=st.sidebar.radio("Selecione:",pg_opts,index=pg_opts.index(st.session_state.pg_sel),key="nav_k")
st.sidebar.divider(); uploaded_file=st.sidebar.file_uploader("Selecione DOCX",type="docx",key="upl_k")
st.session_state.dbg_cb_key=st.sidebar.checkbox("Mostrar Debug Info",value=st.session_state.dbg_cb_key,key="dbg_k")

if uploaded_file:
    if st.session_state.f_name!=uploaded_file.name: 
        with st.spinner("Processando novo documento..."):
            st.session_state.s_gemini,st.session_state.cfg_sugs=[],{}
            st.session_state.f_name=uploaded_file.name
            txt,tbls=extrair_conteudo_docx(uploaded_file);st.session_state.doc_ctx={"texto":txt,"tabelas":tbls}
            if txt or tbls:
                sugs=analisar_documento_com_gemini(txt,tbls);st.session_state.s_gemini=sugs
                st.session_state.cfg_sugs={s.get("id",f"s_{i}_{hash(s.get('titulo'))}"):{"aceito":True,"titulo_editado":s.get("titulo","S/T"),"dados_originais":s} for i,s in enumerate(sugs)}
            else: st.sidebar.warning("Nenhum conteúdo extraído.")
    if st.session_state.dbg_cb_key and (st.session_state.doc_ctx["texto"] or st.session_state.doc_ctx["tabelas"]):
        with st.expander("Debug: Conteúdo DOCX (após extração e tipos)",expanded=False):
            st.text_area("Texto (amostra)",st.session_state.doc_ctx["texto"][:1000],height=80)
            for t_dbg in st.session_state.doc_ctx["tabelas"]:
                st.write(f"ID: {t_dbg['id']}, Nome: {t_dbg['nome']}")
                try: st.dataframe(t_dbg['dataframe'].head().astype(str).fillna("-"))
                except: st.text(f"Head:\n{t_dbg['dataframe'].head().to_string(na_rep='-')}")
                st.write("Tipos:",t_dbg['dataframe'].dtypes.to_dict())
    if st.session_state.s_gemini:
        st.sidebar.divider();st.sidebar.header("⚙️ Configurar Sugestões")
        for sug_sb in st.session_state.s_gemini:
            s_id,cfg_sb=sug_sb['id'],st.session_state.cfg_sugs.get(sug_sb['id'])
            if not cfg_sb: cfg_sb=st.session_state.cfg_sugs[s_id]={"aceito":True,"titulo_editado":sug_sb.get("titulo","S/T"),"dados_originais":sug_sb}
            with st.sidebar.expander(f"{cfg_sb['titulo_editado']}",expanded=False):
                st.caption(f"Tipo: {sug_sb.get('tipo_sugerido')} | Fonte: {sug_sb.get('fonte_id')}")
                cfg_sb["aceito"]=st.checkbox("Incluir?",value=cfg_sb["aceito"],key=f"acc_{s_id}")
                cfg_sb["titulo_editado"]=st.text_input("Título",value=cfg_sb["titulo_editado"],key=f"tit_{s_id}")
else: 
    if st.session_state.pg_sel=="Dashboard Principal": st.info("Upload DOCX na barra lateral.")

if st.session_state.pg_sel=="Dashboard Principal":
    st.title("📊 Dashboard de Insights")
    if uploaded_file and st.session_state.s_gemini:
        kpis,outros=[],[]
        for s_id,s_cfg in st.session_state.cfg_sugs.items():
            if s_cfg["aceito"]: item={"titulo":s_cfg["titulo_editado"],**s_cfg["dados_originais"]};(kpis if item.get("tipo_sugerido")=="kpi" else outros).append(item)
        render_kpis(kpis)
        if st.session_state.dbg_cb_key:
             with st.expander("Debug: Elementos para Dashboard (Não-KPI)",expanded=False): st.json({"Outros":outros},expanded=False)
        if outros:
            cols_d,idx_d,cnt_d=st.columns(2),0,0
            for item_m in outros:
                if item_m.get("tipo_sugerido")=="lista_swot":continue
                with cols_d[idx_d%2]:
                    st.subheader(item_m["titulo"]);df_p,rend=None,False
                    params,tipo,fonte=item_m.get("parametros",{}),item_m.get("tipo_sugerido"),item_m.get("fonte_id")
                    try:
                        if params.get("dados"):df_p=pd.DataFrame(params["dados"])
                        elif str(fonte).startswith("doc_tabela_"):df_p=next((t["dataframe"] for t in st.session_state.doc_ctx["tabelas"] if t["id"]==fonte),None)
                        if tipo=="tabela_dados":
                            df_t=None
                            if str(fonte).startswith("texto_") and params.get("dados"):
                                df_t=pd.DataFrame(params.get("dados"));
                                if params.get("colunas_titulo"):df_t.columns=params.get("colunas_titulo")
                            else:id_t=params.get("id_tabela_original",fonte);df_t=next((t["dataframe"] for t in st.session_state.doc_ctx["tabelas"] if t["id"]==id_t),None)
                            if df_t is not None:try:st.dataframe(df_t.astype(str).fillna("-"))
                                                except:st.text(df_t.to_string(na_rep='-'));rend=True
                            else:st.warning(f"Tabela '{item_m['titulo']}' (Fonte:{fonte}) não encontrada.")
                        elif df_p is not None and tipo in ["grafico_barras","grafico_linha","grafico_dispersao","grafico_pizza"]:
                            if render_plotly_chart(item_m,df_p):rend=True
                        elif tipo=='mapa':st.info(f"Mapa '{item_m['titulo']}' não implementado.");rend=True
                        if not rend and tipo not in ["kpi","lista_swot","mapa"]:st.info(f"'{item_m['titulo']}' ({tipo}) não gerado.")
                    except Exception as e:st.error(f"Erro render '{item_m['titulo']}': {e}")
                if rend:idx_d+=1;cnt_d+=1
            if cnt_d==0 and any(c['aceito'] and c['dados_originais'].get('tipo_sugerido') not in ['kpi','lista_swot'] for c in st.session_state.cfg_sugs.values()):
                st.info("Nenhum gráfico/tabela (além de KPIs/SWOTs) pôde ser gerado.")
        elif not kpis and not uploaded_file:pass
        elif not kpis and not outros and uploaded_file and st.session_state.s_gemini:st.info("Nenhum elemento selecionado/gerado.")
elif st.session_state.pg_sel=="Análise SWOT Detalhada":
    st.title("🔬 Análise SWOT Detalhada")
    if not uploaded_file:st.warning("Upload DOCX na barra lateral.")
    elif not st.session_state.s_gemini:st.info("Aguardando processamento/sugestões.")
    else:
        swot_sugs=[s_cfg["dados_originais"] for s_id,s_cfg in st.session_state.cfg_sugs.items() if s_cfg["aceito"] and s_cfg["dados_originais"].get("tipo_sugerido")=="lista_swot"]
        if not swot_sugs:st.info("Nenhuma análise SWOT sugerida/selecionada.")
        else:
            for swot_item in swot_sugs:
                render_swot_card(swot_item.get("titulo","SWOT"),swot_item.get("parametros",{}),card_key_prefix=swot_item.get("id","swot"))
if uploaded_file is None and st.session_state.f_name is not None:
    keys_to_preserve=["debug_cb_widget_key_vfinal_corrected_again_swnan","uploader_main_key_vfinal_corrected_again_swnan","nav_radio_key"] # Chaves dos widgets principais
    current_keys=list(st.session_state.keys())
    for k_cl in current_keys:
        is_widget_key_to_preserve = any(k_cl.startswith(preserved_prefix) for preserved_prefix in ["acc_s_","tit_s_","px_s_","py_s_","pcat_s_","pval_s_"])
        if k_cl not in keys_to_preserve and not is_widget_key_to_preserve:
            del st.session_state[k_cl]
    for k_r,dv_r in [("s_gemini",[]),("cfg_sugs",{}),("doc_ctx",{"texto":"","tabelas":[]}),("f_name",None),("debug_checkbox_key_main",False),("pg_sel","Dashboard Principal")]:st.session_state.setdefault(k_r,dv_r)
    st.experimental_rerun()